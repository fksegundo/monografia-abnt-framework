# -*- coding: utf-8 -*-
"""Insere figuras/imagens num documento .docx, substituindo placeholders.

Procura por placeholders de texto (ex.: "INSERIR FIGURA 1 AQUI") e os
substitui por uma imagem centralizada, com legenda acima e fonte abaixo.
Os mapeamentos podem ser fornecidos via JSON (--config) ou inline. Formato
esperado do JSON:

    {
      "INSERIR FIGURA 1 AQUI": {
        "imagem": "assets/figura1.png",
        "legenda": "Figura 1 — Descrição.",
        "largura_cm": 14,
        "fonte": "Fonte: Elaborado pelo autor (2026)."
      }
    }

Uso:
    python -m scripts.inserir_figuras <arquivo.docx> --config mapeamento.json [--saida saida.docx]
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


def estilizar_legenda_ou_fonte(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)


def inserir_figuras(doc_path: Path, config: dict, saida: Path | None = None) -> None:
    if not doc_path.exists():
        raise SystemExit(f"Arquivo não encontrado: {doc_path}")

    # Resolve caminhos relativos ao diretório raiz do framework/arquivo de config
    base = Path(__file__).resolve().parents[1]

    print(f"Abrindo {doc_path} para inserção de imagens...")
    doc = Document(str(doc_path))
    saida = saida or doc_path

    inserted = []
    for p in doc.paragraphs:
        for placeholder, spec in config.items():
            if placeholder in p.text:
                img_rel = spec.get("imagem", "")
                img_path = Path(img_rel)
                if not img_path.is_absolute():
                    img_path = base / img_path
                if not img_path.exists():
                    print(f"  [SKIP] {placeholder} — arquivo não encontrado: {img_path}")
                    break

                legenda = spec.get("legenda", "")
                largura = Cm(spec.get("largura_cm", 14))
                fonte = spec.get("fonte", "Fonte: Elaborado pelo autor.")

                if legenda:
                    cap = p.insert_paragraph_before(legenda)
                    estilizar_legenda_ou_fonte(cap)

                p.text = ""
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(str(img_path), width=largura)

                fonte_p = _insert_paragraph_after(p, fonte)
                estilizar_legenda_ou_fonte(fonte_p)

                inserted.append(placeholder)
                print(f"  [OK] {placeholder} -> {img_path.name}")
                break

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(saida))
    print(f"\nProcesso concluído. {len(inserted)} imagens inseridas em {saida}.")


def _insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def main() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass
    parser = argparse.ArgumentParser(description="Insere figuras em .docx por placeholder.")
    parser.add_argument("arquivo", help="Caminho do arquivo .docx.")
    parser.add_argument("--config", help="Arquivo JSON com o mapeamento placeholder -> imagem.")
    parser.add_argument("--saida", help="Caminho de saída (se omitido, modifica in-place).")
    args = parser.parse_args()

    if not args.config:
        raise SystemExit("Forneça --config com o mapeamento placeholder -> imagem.")

    with open(args.config, "r", encoding="utf-8-sig") as f:
        config = json.load(f)

    inserir_figuras(Path(args.arquivo), config, Path(args.saida) if args.saida else None)


if __name__ == "__main__":
    main()
