# -*- coding: utf-8 -*-
"""Gerador de monografia a partir de drafts em Markdown.

Lê os drafts em `drafts/`, injeta seu conteúdo em um template .docx (ABNT)
e gera o documento final formatado (com sumário e capturas de títulos)
em `output/`.

Uso:
    python -m scripts.gerar_monografia [--template template/modelo.docx] \
        [--saida output/monografia.docx] [--drafts drafts]

Pré-requisitos:
    - O template .docx deve conter placeholders "Espaço do resumo.",
      "Versão traduzida do resumo," e os títulos usados pelos drafts.
    - Os drafts devem ser .md numerados sequencialmente (ex.: 00_pre_textual,
      01_introducao, 02_fundamentacao, ...).
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

try:
    from formatacao.abnt import PRETO, VERMELHO, add_markdown_runs, estilo_codigo, estilo_legenda_ou_fonte, estilo_titulo
except ImportError:
    from scripts.formatacao.abnt import PRETO, VERMELHO, add_markdown_runs, estilo_codigo, estilo_legenda_ou_fonte, estilo_titulo

RAIZ = Path(__file__).resolve().parents[1]


def limpar_sujeira_markdown(texto):
    """Remove os '#' de título markdown e retorna (texto, nivel)."""
    text = texto.strip()
    if not text.startswith("#"):
        return text, None
    level = text.split(" ")[0].count("#") - 1
    clean_text = text.lstrip("#").strip()
    return clean_text, level


def parse_markdown_blocks(content):
    """Divide o markdown em blocos de texto/código/tabela."""
    blocks = []
    paragraph = []
    code = []
    in_code = False

    for line in content.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(code)))
                code = []
                in_code = False
            else:
                if paragraph:
                    blocks.append(("text", "\n".join(paragraph).strip()))
                    paragraph = []
                in_code = True
            continue

        if in_code:
            code.append(stripped)
            continue

        if not stripped:
            if paragraph:
                blocks.append(("text", "\n".join(paragraph).strip()))
                paragraph = []
            continue

        paragraph.append(stripped)

    if code:
        blocks.append(("code", "\n".join(code)))
    if paragraph:
        blocks.append(("text", "\n".join(paragraph).strip()))
    return blocks


def is_markdown_table(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return (
        len(lines) >= 2
        and all(line.startswith("|") and line.endswith("|") for line in lines)
        and set(lines[1].replace("|", "").replace(":", "").replace("-", "").strip()) == set()
    )


def add_markdown_table(doc, text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    rows = []
    for idx, line in enumerate(lines):
        if idx == 1:
            continue
        rows.append([cell.strip() for cell in line.strip("|").split("|")])
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            add_markdown_runs(p, cell_text, allow_english_italics=True)
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = docx.shared.Pt(10)
                if row_idx == 0:
                    run.bold = True


def extract_pretext_sections(text):
    """Extrai as seções pré-textuais (LISTA DE ...) dos blocos markdown."""
    sections = {}
    current = None
    items = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            if current:
                sections[current] = "\n".join(items).strip()
            title = stripped[2:].strip()
            current = title if title.startswith("LISTA DE ") else None
            items = []
            continue
        if current and stripped:
            items.append(stripped)
    if current:
        sections[current] = "\n".join(items).strip()
    return sections


def extrair_resumo_abstract(text0):
    """Extrai o texto de RESUMO, ABSTRACT e lista de siglas de um draft pré-textual."""
    str_resumo = ""
    str_abstract = ""
    str_siglas = ""
    current = None
    for bloco in text0.split("\n\n"):
        if "RESUMO" in bloco and "# " in bloco:
            current = "R"
            continue
        if "ABSTRACT" in bloco and "# " in bloco:
            current = "A"
            continue
        if "LISTA DE ABR" in bloco and "# " in bloco:
            current = "S"
            continue
        if "LISTA DE FIGURAS" in bloco or "LISTA DE GRÁFICOS" in bloco or "LISTA DE QUADROS" in bloco or "LISTA DE TABELAS" in bloco:
            if "# " in bloco:
                current = None
                continue
        if current == "R":
            str_resumo += bloco + "\n\n"
        elif current == "A":
            str_abstract += bloco + "\n\n"
        elif current == "S":
            str_siglas += bloco + "\n"
    return str_resumo, str_abstract, str_siglas


def preparar_pretextuais(doc, pretext, pretext_list_map):
    """Substitui placeholders pré-textuais do template pelo conteúdo dos drafts."""
    str_resumo, str_abstract, str_siglas = extrair_resumo_abstract(pretext)

    siglas_trash = ["BNT\tAssocia", "BUCBP\t", "CRA\t", "CRC\t", "NBR\t", "NPJ\t", "OAB\t", "UNDB\t", "TCC\tTra"]
    inserted_siglas = False
    current_list_title = None

    for p in doc.paragraphs:
        if p.text.strip() in pretext_list_map:
            current_list_title = p.text.strip()
            continue

        if "Espaço do resumo." in p.text:
            p.clear()
            add_markdown_runs(p, str_resumo.strip(), allow_english_italics=True)
            estilo_titulo(p)

        if "Versão traduzida do resumo," in p.text:
            p.clear()
            add_markdown_runs(p, str_abstract.strip(), allow_english_italics=False)
            estilo_titulo(p)

        if any(st in p.text for st in siglas_trash):
            if not inserted_siglas:
                p.text = str_siglas.strip()
                estilo_titulo(p)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.first_line_indent = None
                inserted_siglas = True
            else:
                p._element.getparent().remove(p._element)

        if "[Fonte 12, Esp" in p.text:
            p._element.getparent().remove(p._element)

        if "Deve ser sempre atualizado" in p.text and current_list_title:
            p.text = pretext_list_map.get(current_list_title, "")
            estilo_titulo(p)
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            current_list_title = None


def injetar_draft(doc, path, in_references=False):
    """Injecta o conteúdo de um draft .md no documento."""
    content = Path(path).read_text(encoding="utf-8-sig")

    for block_type, bloco in parse_markdown_blocks(content):
        bloco = bloco.strip()
        if not bloco:
            continue

        if block_type == "text" and is_markdown_table(bloco):
            add_markdown_table(doc, bloco)
            continue

        novo_p = doc.add_paragraph()

        if block_type == "code":
            for idx, line in enumerate(bloco.splitlines()):
                if idx:
                    novo_p.add_run().add_break()
                run = novo_p.add_run(line)
                run.font.name = "Courier New"
                run.font.size = docx.shared.Pt(10)
                run.font.color.rgb = PRETO
            estilo_codigo(novo_p)
            continue

        if "[INSERIR " in bloco:
            run = novo_p.add_run(bloco)
            run.bold = True
            run.font.color.rgb = VERMELHO
            run.font.name = "Arial"
            run.font.size = docx.shared.Pt(12)
            novo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if re.match(r"^(Quadro|Figura|Gráfico|Grafico) \d+\s+[—-]", bloco) or bloco.startswith("Fonte:"):
            add_markdown_runs(novo_p, bloco, allow_english_italics=False)
            estilo_legenda_ou_fonte(novo_p)
            continue

        texto_limpo, nivel = limpar_sujeira_markdown(bloco)

        if nivel is not None:
            if nivel == 0 and len(doc.paragraphs) > 5:
                run = novo_p.add_run()
                run.add_break(WD_BREAK.PAGE)
            add_markdown_runs(novo_p, texto_limpo, allow_english_italics=False)
            estilo_titulo(novo_p, nivel=nivel)
            if texto_limpo == "REFERÊNCIAS":
                novo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                in_references = True
        else:
            add_markdown_runs(novo_p, texto_limpo, allow_english_italics=not in_references)
            estilo_titulo(novo_p)
            if in_references:
                novo_p.paragraph_format.first_line_indent = None
                novo_p.paragraph_format.line_spacing = 1.0
                novo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return in_references


def main() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass
    parser = argparse.ArgumentParser(description="Gera monografia ABNT a partir de drafts Markdown.")
    parser.add_argument("--template", default=str(RAIZ / "template" / "modelo.docx"), help="Template .docx de origem.")
    parser.add_argument("--saida", default=str(RAIZ / "output" / "monografia.docx"), help="Caminho do documento gerado.")
    parser.add_argument("--md-saida", default=str(RAIZ / "output" / "monografia.md"), help="Caminho do markdown concatenado.")
    parser.add_argument("--drafts", default=str(RAIZ / "drafts"), help="Pasta contendo os drafts .md.")
    parser.add_argument("--pre-textual", help="Draft pré-textual (resumo/abstract/siglas/listas). Se omitido, usa o 1o arquivo ordenado da pasta.")
    args = parser.parse_args()

    template = Path(args.template)
    drafts_dir = Path(args.drafts)

    if not template.exists():
        raise SystemExit(f"Template não encontrado: {template}")

    print("--- Gerando documento ---")
    print(f"Template: {template}")
    doc = docx.Document(str(template))

    # Ordena drafts numericamente (00_, 01_, ...)
    draft_files = sorted(
        [p for p in drafts_dir.glob("*.md")],
        key=lambda p: p.name,
    )
    if not draft_files:
        raise SystemExit(f"Nenhum draft .md encontrado em {drafts_dir}")

    print(f"Encontrados {len(draft_files)} drafts.")

    # PASSO 1: localizar os marcos de INTRODUÇÃO/DESENVOLVIMENTO no template
    idx_intro = -1
    idx_desenv = -1
    for i, p in enumerate(doc.paragraphs):
        if "1 INTRODUÇÃO" in p.text or "1 INTRODUCAO" in p.text.upper():
            idx_intro = i
        if "2 DESENVOLVIMENTO" in p.text or "2 DESENVOLVIMENTO" in p.text.upper():
            idx_desenv = i
            break

    # PASSO 2: apagar o conteúdo placeholder entre INTRODUÇÃO e DESENVOLVIMENTO
    if idx_intro != -1 and idx_desenv != -1:
        for i in range(len(doc.paragraphs) - 1, idx_intro, -1):
            try:
                doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
            except Exception:
                pass

    if idx_intro != -1:
        doc.paragraphs[idx_intro]._element.getparent().remove(doc.paragraphs[idx_intro]._element)

    # PASSO 3: definir o draft pré-textual (resumo/abstract)
    pretext_path = None
    pretext_list_map = {
        "LISTA DE FIGURAS": "",
        "LISTA DE GRÁFICOS": "",
        "LISTA DE QUADROS": "",
        "LISTA DE TABELAS": "",
    }
    if args.pre_textual:
        pretext_path = Path(args.pre_textual)
    else:
        # O primeiro draft ordenado é considerado pré-textual
        pretext_path = draft_files[0]

    if pretext_path and pretext_path.exists():
        pretext = Path(pretext_path).read_text(encoding="utf-8-sig")
        sections = extract_pretext_sections(pretext)
        for k in pretext_list_map:
            pretext_list_map[k] = sections.get(k, "")
        preparar_pretextuais(doc, pretext, pretext_list_map)
        draft_files = [p for p in draft_files if p != pretext_path]

    # PASSO 4: injetar os drafts do corpo
    print("Injetando conteúdo dos drafts...")
    in_refs = False
    for draft in draft_files:
        print(f"  -> {draft.name}")
        in_refs = injetar_draft(doc, draft, in_refs)

    # PASSO 5: salvar
    saida_path = Path(args.saida)
    saida_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(saida_path))
    print(f"Documento gerado: {saida_path}")

    md_saida = Path(args.md_saida)
    if md_saida:
        md_saida.parent.mkdir(parents=True, exist_ok=True)
        with open(str(md_saida), "w", encoding="utf-8") as f:
            if pretext_path and pretext_path.exists():
                f.write(Path(pretext_path).read_text(encoding="utf-8-sig").strip() + "\n\n")
            for draft in draft_files:
                f.write(Path(draft).read_text(encoding="utf-8-sig").strip() + "\n\n")
        print(f"Markdown concatenado: {md_saida}")


if __name__ == "__main__":
    main()
