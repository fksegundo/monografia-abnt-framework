# -*- coding: utf-8 -*-
"""Gera o template .docx ABNT genérico e anônimo usado pelo framework.

O template é a base pré-textual (capa, folha de rosto, listas, resumo/abstract)
que o orquestrador `gerar_monografia.py` usa como moldura. O conteúdo textual
(INTRODUÇÃO em diante) é apagado e reinjetado a partir dos drafts, por isso o
template precisa apenas:

- dos elementos pré-textuais (capa, folha de rosto, resumo, listas);
- dos placeholders exatos que o orquestrador substitui:
    * "Espaço do resumo."        (no corpo do RESUMO);
    * "Versão traduzida do resumo," (no corpo do ABSTRACT);
    * títulos exatos das listas: "LISTA DE FIGURAS", "LISTA DE GRÁFICOS",
      "LISTA DE QUADROS", "LISTA DE TABELAS", seguidos de uma linha
      "Deve ser sempre atualizado..." para o script preencher;
    * os marcos de corpo "1 INTRODUÇÃO" e "2 DESENVOLVIMENTO".

Nenhum dado pessoal ou de instituição é incluído: instituição, curso, cidade,
área e orientação entram como placeholders genéricos que o usuário preenche no
template (ex.: "[INSTITUIÇÃO]", "[CURSO]", "[CIDADE]", "NOME COMPLETO",
"[ÁREA]", "Nome do Orientador").

Uso:
    python -m scripts.criar_template [--saida template/modelo.docx]
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parents[1]

FONTE = "Arial"
PRETO = RGBColor(0, 0, 0)


def _reconfigurar_stdout() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def _par(doc, text="", size=12, bold=None, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_after=0, space_before=0, line=1.5, indent=None, center_text=False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_text else align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if indent is not None:
        pf.first_line_indent = indent
    if text:
        run = p.add_run(text)
        run.font.name = FONTE
        run.font.size = Pt(size)
        run.font.color.rgb = PRETO
        run.bold = bold
        run.italic = italic
    return p


def _quebra(doc):
    doc.add_page_break()


def criar(saida: Path) -> Path:
    doc = Document()

    # Formato A4 e margens ABNT: esquerda/superior 3 cm, direita/inferior 2 cm
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.top_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # Estilo de parágrafo padrão
    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(12)

    # ---------------- CAPA ----------------
    _par(doc, space_before=60)
    _par(doc, "[INSTITUIÇÃO - CENTRO UNIVERSITÁRIO / FACULDADE]", center_text=True, bold=True)
    _par(doc, "[CURSO - EX.: ENGENHARIA DE SOFTWARE]", center_text=True, bold=True)
    _par(doc, space_after=90)
    _par(doc, "NOME COMPLETO [Fonte 12, negrito]", center_text=True)
    _par(doc, "TÍTULO [Fonte 12, negrito, caixa alta]: subtítulo se houver [Fonte 12, sem negrito]", center_text=True)
    _par(doc, space_before=90)
    _par(doc, "[CIDADE]", center_text=True)
    _par(doc, "2026", center_text=True)
    _quebra(doc)

    # ---------------- FOLHA DE ROSTO ----------------
    _par(doc, "NOME COMPLETO", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "TÍTULO: subtítulo se houver", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, space_after=60)
    _par(doc, "[TEXTO DA FOLHA DE ROSTO conforme as normas da instituição. Ex.: "
              "Monografia apresentada ao Curso de [NOME DO CURSO] da [INSTITUIÇÃO] como requisito "
              "parcial para obtenção do grau de [GRAU] em [ÁREA].]", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5)
    _par(doc, "Orientador: Prof. Me. Nome do Orientador.", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, space_before=90)
    _par(doc, "[CIDADE]", center_text=True)
    _par(doc, "2026", center_text=True)
    _quebra(doc)

    # ---------------- FICHA CATALOGRÁFICA ----------------
    _par(doc, "FICHA CATALOGRÁFICA", center_text=True, bold=True)
    _par(doc, "NOME COMPLETO", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "TÍTULO: subtítulo se houver", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, space_after=60)
    _par(doc, "[TEXTO DA FICHA CATALOGRÁFICA conforme a biblioteca da instituição.]",
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5)
    _par(doc, "Aprovada em: _____/____/_____.", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "BANCA EXAMINADORA:", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "_____________________________________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "Prof. Me. Nome Completo (Orientador)", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "Mestre em [ÁREA]", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "[INSTITUIÇÃO]", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "_____________________________________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "Prof. Me. Nome Completo", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "Mestre em [ÁREA]", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "[INSTITUIÇÃO]", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "_____________________________________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "Prof. Me. Nome Completo", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "Mestre em [ÁREA]", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "[INSTITUIÇÃO]", align=WD_ALIGN_PARAGRAPH.LEFT)
    _quebra(doc)

    # ---------------- DEDICATÓRIA ----------------
    _par(doc, space_before=120)
    _par(doc, "Dedico este trabalho a minha família.", align=WD_ALIGN_PARAGRAPH.RIGHT, line=1.5)
    _quebra(doc)

    # ---------------- AGRADECIMENTOS ----------------
    _par(doc, "AGRADECIMENTOS", center_text=True, bold=True)
    _par(doc, "Espaço destinado aos agradecimentos. Embora seja de formatação livre, costuma-se "
              "utilizar fonte 12, espaçamento 1,5 mm, recuo de 1,25 cm ou 2 cm.", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5)
    _quebra(doc)

    # ---------------- EPÍGRAFE ----------------
    _par(doc, space_before=120)
    _par(doc, "“Espaço destinado à epígrafe.”", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)
    _par(doc, "(SOBRENOME, ano, p. 1).", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _quebra(doc)

    # ---------------- RESUMO ----------------
    _par(doc, "RESUMO", center_text=True, bold=True)
    _par(doc, "Espaço do resumo. Deve seguir as orientações para elaboração de resumo informativo "
              "na NBR 6028.", align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5)
    _par(doc, "Palavras-chave: Palavra 1. Palavra 2. Palavra 3. Palavra 4. Palavra 5.",
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5)
    _quebra(doc)

    # ---------------- ABSTRACT ----------------
    _par(doc, "ABSTRACT", center_text=True, bold=True)
    _par(doc, "Versão traduzida do resumo, seguindo a mesma formatação.",
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5)
    _par(doc, "Keywords: Keyword 1. Keyword 2. Keyword 3. Keyword 4. Keyword 5.",
         align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5)
    _quebra(doc)

    # ---------------- LISTAS (placeholders de preenchimento) ----------------
    for titulo in ["LISTA DE FIGURAS", "LISTA DE GRÁFICOS", "LISTA DE QUADROS", "LISTA DE TABELAS"]:
        _par(doc, titulo, center_text=True, bold=True)
        _par(doc, "[Item de exemplo da lista.]", align=WD_ALIGN_PARAGRAPH.LEFT)
        _par(doc, "Deve ser sempre atualizado, à medida que novos elementos forem inseridos no texto.",
             align=WD_ALIGN_PARAGRAPH.LEFT)
        _quebra(doc)

    # ---------------- LISTA DE ABREVIATURAS E SIGLAS ----------------
    _par(doc, "LISTA DE ABREVIATURAS E SIGLAS", center_text=True, bold=True)
    _par(doc, "TCC — Trabalho de Conclusão de Curso", align=WD_ALIGN_PARAGRAPH.LEFT)
    _quebra(doc)

    # ---------------- SUMÁRIO ----------------
    _par(doc, "SUMÁRIO", center_text=True, bold=True)
    _par(doc, "1 INTRODUÇÃO", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "2 DESENVOLVIMENTO", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "3 TÓPICO PRIMÁRIO DO DESENVOLVIMENTO", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "4 RESULTADOS E DISCUSSÕES", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "5 CONSIDERAÇÕES FINAIS", align=WD_ALIGN_PARAGRAPH.LEFT)
    _par(doc, "REFERÊNCIAS", align=WD_ALIGN_PARAGRAPH.LEFT)
    _quebra(doc)

    # ---------------- MARCOS DE CORPO (apagados e reinjetados pelo orquestrador) ----------------
    _par(doc, "1 INTRODUÇÃO", center_text=True, bold=True)
    _par(doc, "2 DESENVOLVIMENTO", center_text=True, bold=True)

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(saida))
    return saida


def main() -> None:
    _reconfigurar_stdout()
    parser = argparse.ArgumentParser(description="Gera o template ABNT genérico e anônimo.")
    parser.add_argument("--saida", default=str(RAIZ / "template" / "modelo.docx"), help="Caminho do .docx de saída.")
    args = parser.parse_args()
    out = criar(Path(args.saida))
    print(f"Template gerado: {out}")


if __name__ == "__main__":
    main()
