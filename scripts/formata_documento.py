# -*- coding: utf-8 -*-
"""Formata um documento .docx existente segundo normas ABNT.

Aplica fontes (Arial), recuos, espaçamento, tratamentos de títulos
(Heading 1-5) e adiciona um sumário automático (campo TOC) após o título
"SUMÁRIO".

Uso:
    python -m scripts.formata_documento <arquivo.docx> [--saida saida.docx]

Se --saida for omitido, o arquivo é formatado in-place (com backup .backup).
"""

from __future__ import annotations

import argparse
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


HEADING_RE = re.compile(r"^(\d+(?:\.\d+){0,4})\s+\S+")
PRETO = RGBColor(0, 0, 0)


def set_run_font(run, size=12, bold=None):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = PRETO
    if bold is not None:
        run.bold = bold


def set_style_font(style, size=12, bold=False):
    font = style.font
    font.name = "Arial"
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Arial")
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = PRETO


def get_heading_level(text: str):
    stripped = text.strip()
    if stripped == "REFERÊNCIAS":
        return 1
    match = HEADING_RE.match(stripped)
    if not match:
        return None
    prefix = match.group(1)
    if "." not in prefix:
        return 1
    return min(prefix.count(".") + 1, 5)


def set_paragraph_outline_level(paragraph, level: int):
    ppr = paragraph._p.get_or_add_pPr()
    for old in ppr.findall(qn("w:outlineLvl")):
        ppr.remove(old)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level - 1))
    ppr.append(outline)


def style_heading(paragraph, level: int):
    for style_name in (f"Heading {level}", f"Título {level}"):
        try:
            paragraph.style = style_name
            break
        except KeyError:
            continue
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_outline_level(paragraph, level)
    for run in paragraph.runs:
        set_run_font(run, size=12, bold=True)


def style_body(paragraph):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=12)


def style_simple_center(paragraph, bold=True):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=12, bold=bold)


def style_references_body(paragraph):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=12)


def add_toc_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    toc_p = Paragraph(new_p, paragraph._parent)

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-5" \h \z \u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Clique com o botão direito e escolha Atualizar Campo para gerar o sumário automático."
    run.append(text)
    fld.append(run)
    new_p.append(fld)

    toc_p.paragraph_format.first_line_indent = None
    toc_p.paragraph_format.line_spacing = 1.0
    toc_p.paragraph_format.space_after = Pt(12)
    for run_obj in toc_p.runs:
        set_run_font(run_obj, size=12)


def enable_update_fields(doc: Document):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def remove_existing_toc_hint(paragraphs):
    hints = [
        "Clique com o botão direito e escolha Atualizar Campo",
        "No table of contents entries found",
    ]
    for paragraph in list(paragraphs):
        if any(hint in paragraph.text for hint in hints):
            paragraph._element.getparent().remove(paragraph._element)


def formatar_abnt(arquivo: Path, saida: Path | None = None) -> None:
    if not arquivo.exists():
        raise SystemExit(f"Arquivo não encontrado: {arquivo}")

    backup = None
    output_path = saida or arquivo

    # Se for formatação in-place, faz backup
    if saida is None:
        backup = arquivo.with_suffix(arquivo.suffix + ".backup")
        if not backup.exists():
            shutil.copy2(arquivo, backup)

    doc = Document(str(arquivo))

    for style_name in [
        "Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5",
        "Título 1", "Título 2", "Título 3", "Título 4", "Título 5",
    ]:
        try:
            set_style_font(doc.styles[style_name], bold=style_name.startswith(("Heading", "Título")))
        except KeyError:
            pass

    remove_existing_toc_hint(doc.paragraphs)

    in_references = False
    toc_paragraph = None
    changed_headings = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            paragraph.style = "Normal"
            continue

        if text == "SUMÁRIO":
            toc_paragraph = paragraph
            style_simple_center(paragraph, bold=True)
            in_references = False
            continue

        level = get_heading_level(text)
        if level is not None:
            style_heading(paragraph, level)
            changed_headings += 1
            in_references = text == "REFERÊNCIAS"
            continue

        if text in {"RESUMO", "ABSTRACT", "LISTA DE FIGURAS", "LISTA DE GRÁFICOS",
                    "LISTA DE QUADROS", "LISTA DE ABREVIATURAS E SIGLAS"}:
            style_simple_center(paragraph, bold=True)
            continue

        if in_references:
            style_references_body(paragraph)
        else:
            style_body(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = None
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_run_font(run, size=10)

    if toc_paragraph is not None:
        add_toc_after(toc_paragraph)

    enable_update_fields(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"Arquivo formatado: {output_path}")
    print(f"Títulos ajustados para Heading 1-5: {changed_headings}")
    if backup:
        print(f"Backup: {backup}")


def main() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass
    parser = argparse.ArgumentParser(description="Formata documento .docx segundo ABNT.")
    parser.add_argument("arquivo", help="Caminho do arquivo .docx.")
    parser.add_argument("--saida", help="Caminho de saída (se omitido, formata in-place com backup).")
    args = parser.parse_args()
    formatar_abnt(Path(args.arquivo), Path(args.saida) if args.saida else None)


if __name__ == "__main__":
    main()
