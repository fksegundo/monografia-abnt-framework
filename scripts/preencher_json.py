# -*- coding: utf-8 -*-
"""Preenche um documento .docx com conteúdo vindo de um arquivo JSON.

O JSON mapeia o prefixo de um cabeçalho (ex.: "1", "2.1", "REF") para uma
string com os parágrafos que devem ser inseridos logo abaixo dele. Os
parágrafos são separados por linhas em branco (\n\n) no valor.

Uso:
    python -m scripts.preencher_json <arquivo.docx> <conteudo.json> <saida.docx>
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def get_prefix(text: str):
    if text.strip().upper().startswith("REFER"):
        return "REF"
    m = re.match(r"^(\d+(?:\.\d+)*)\s+", text.strip())
    if m:
        return m.group(1)
    if text.strip().upper().startswith("1 INTRO"):
        return "1"
    return None


def fill_doc(doc_path: Path, json_path: Path, out_path: Path) -> None:
    print("Carregando conteúdos...")
    with open(str(json_path), "r", encoding="utf-8-sig") as f:
        data = json.load(f)

    print(f"Abrindo arquivo {doc_path}...")
    doc = Document(str(doc_path))

    for prefix, texto_paragrafos in data.items():
        next_p = None
        found_target = False

        for p in doc.paragraphs:
            if not found_target:
                if get_prefix(p.text) == prefix:
                    found_target = True
            else:
                next_p = p
                break

        if found_target:
            textos = texto_paragrafos.split("\n\n")
            for txt in textos:
                if not txt.strip():
                    continue
                if next_p:
                    new_p = next_p.insert_paragraph_before()
                else:
                    new_p = doc.add_paragraph()

                run = new_p.add_run(txt.strip())
                run.font.name = "Arial"
                run.font.size = Pt(12)
                new_p.paragraph_format.first_line_indent = Pt(35.4)  # ~1.25cm
                new_p.paragraph_format.space_before = Pt(6)
                new_p.paragraph_format.line_spacing = 1.5
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Preenchimento realizado com sucesso em {out_path}!")


def main() -> None:
    if len(sys.argv) < 4:
        print("Uso: python -m scripts.preencher_json <arquivo.docx> <conteudo.json> <saida.docx>")
        sys.exit(1)
    fill_doc(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))


if __name__ == "__main__":
    main()
