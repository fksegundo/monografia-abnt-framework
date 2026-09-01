# -*- coding: utf-8 -*-
"""Utilitários de formatação ABNT compartilhados entre os scripts.

Centraliza a lógica de formatação (fontes, recuos, espaçamento, negrito,
itálico) reutilizada pelos scripts de geração e formatação de documentos.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONTE_PADRAO = "Arial"
FONTE_CODIGO = "Courier New"
TAMANHO_PADRAO = Pt(12)
TAMANHO_CODIGO = Pt(10)
TAMANHO_LEGENDA = Pt(10)
PRETO = RGBColor(0, 0, 0)
VERMELHO = RGBColor(255, 0, 0)

# Localização padrão do arquivo de termos extras (um por linha; `#` = comentário).
RAIZ = Path(__file__).resolve().parents[2]
DEFAULT_TERMOS_PATH = RAIZ / "config" / "termos_ingles.txt"

# Termos estrangeiros que ficam em itálico no corpo do texto (ABNT exige itálico
# para palavras estrangeiras). Esta é a LISTA BASE genérica, válida para qualquer
# trabalho. Para termos específicos do tema, adicione-os em config/termos_ingles.txt
# (um por linha) — eles são somados a esta lista automaticamente. Preferências:
#   * no config vale usar o termo em caixa/acentuação como aparece no texto (a
#     comparação é case-insensitive);
#   * termos compostos com espaços/hífens funcionam normalmente.
TERMOS_INGLES_BASE = [
    "ad hoc",
    "batch",
    "big data",
    "cloud",
    "cloud computing",
    "dataset",
    "feedback",
    "insight",
    "know-how",
    "machine learning",
    "on-demand",
    "on-premise",
    "open-source",
    "software as a service",
    "streaming",
    "workflow",
]


def ler_termos_ingles(arquivo: str | Path | None = None) -> list[str]:
    """Lê os termos extras de um arquivo (um por linha; ignora vazias e `#`).

    Se o arquivo não existir, retorna lista vazia.
    """
    caminho = Path(arquivo) if arquivo else DEFAULT_TERMOS_PATH
    termos: list[str] = []
    try:
        for linha in caminho.read_text(encoding="utf-8-sig").splitlines():
            linha = linha.strip()
            if linha and not linha.startswith("#"):
                termos.append(linha)
    except FileNotFoundError:
        pass
    return termos


def todos_termos_ingles(arquivo: str | Path | None = None) -> list[str]:
    """Lista base + termos extras do arquivo de configuração."""
    return TERMOS_INGLES_BASE + ler_termos_ingles(arquivo)


TERMOS_INGLES = todos_termos_ingles()

TERMO_RE = re.compile(
    r"(?<![\w/])("
    + "|".join(re.escape(t) for t in sorted(TERMOS_INGLES, key=len, reverse=True))
    + r")(?![\w/])",
    flags=re.IGNORECASE,
)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+\*|`[^`\n]+`)")


def set_run_font(run, size: int = 12, bold: bool | None = None, color: RGBColor = PRETO) -> None:
    """Aplica fonte/estilo a um run, incluindo a fixação das fontes no XML."""
    run.font.name = FONTE_PADRAO
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONTE_PADRAO)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_run(p, text: str, bold: bool = False, italic: bool = False, code: bool = False) -> None:
    """Adiciona um run com a fonte correta (código em Courier, resto em Arial)."""
    if not text:
        return
    run = p.add_run(text)
    run.font.name = FONTE_CODIGO if code else FONTE_PADRAO
    run.font.size = TAMANHO_CODIGO if code else TAMANHO_PADRAO
    run.font.color.rgb = PRETO
    run.bold = bold
    run.italic = italic


def add_plain_with_english_italics(p, text: str, bold: bool = False, allow_english_italics: bool = True) -> None:
    """Adiciona texto deixando termos estrangeiros em itálico."""
    if not allow_english_italics:
        add_run(p, text, bold=bold)
        return
    pos = 0
    for match in TERMO_RE.finditer(text):
        add_run(p, text[pos:match.start()], bold=bold)
        add_run(p, match.group(0), bold=bold, italic=True)
        pos = match.end()
    add_run(p, text[pos:], bold=bold)


def add_markdown_runs(p, text: str, allow_english_italics: bool = True) -> None:
    """Converte formatação markdown inline (**negrito**, *itálico*, `código`)."""
    pos = 0
    for match in INLINE_RE.finditer(text):
        add_plain_with_english_italics(p, text[pos:match.start()], allow_english_italics=allow_english_italics)
        token = match.group(0)
        if token.startswith("**"):
            add_plain_with_english_italics(p, token[2:-2].replace("*", ""), bold=True, allow_english_italics=allow_english_italics)
        elif token.startswith("*"):
            add_run(p, token[1:-1], italic=True)
        elif token.startswith("`"):
            add_run(p, token[1:-1], code=True)
        pos = match.end()
    add_plain_with_english_italics(p, text[pos:], allow_english_italics=allow_english_italics)


def estilo_titulo(p, nivel: int | None = None) -> None:
    """Aplica estilo ABNT a um parágrafo (título ou corpo)."""
    if nivel is not None:
        try:
            p.style = f"Heading {nivel + 1}"
        except Exception:
            p.style = "Normal"

        ppr = p._p.get_or_add_pPr()
        for old in ppr.findall(qn("w:outlineLvl")):
            ppr.remove(old)
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(nivel))
        ppr.append(outline)

        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            set_run_font(run, bold=True)
    else:
        p.style = "Normal"
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            if run.font.name != FONTE_CODIGO:
                set_run_font(run)


def estilo_codigo(p) -> None:
    """Aplica estilo ABNT a um parágrafo de código."""
    p.style = "Normal"
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = FONTE_CODIGO
        run.font.size = TAMANHO_CODIGO
        run.font.color.rgb = PRETO


def estilo_legenda_ou_fonte(p) -> None:
    """Aplica estilo a legendas de figuras/quadros e às linhas 'Fonte:'."""
    p.style = "Normal"
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ident = re.match(r"^(Quadro|Figura|Gráfico|Grafico) \d+", p.text)
    for run in p.runs:
        set_run_font(run, size=10, bold=bool(ident and run.text.startswith(ident.group(0))))
