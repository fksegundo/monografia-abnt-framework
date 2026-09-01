# -*- coding: utf-8 -*-
"""Substitui/insere tópicos (cabeçalhos) num documento .docx a partir de um
arquivo de texto com a estrutura de tópicos desejada.

Substitui placeholder conhecidos (ex.: "1 INTRODUÇÃO", "2 DESENVOLVIMENTO",
"REFERÊNCIAS") e insere novos tópicos no local lógico, aplicando níveis de
Heading (1-5) e formatação ABNT.

Uso:
    python -m scripts.gerenciador_topicos <arquivo.docx> <topicos.txt> <saida.docx>
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


PRETO = RGBColor(0, 0, 0)


def limpar_texto_topico(linha: str) -> str:
    linha = linha.replace("[", "").replace("]", "").strip()
    linha = re.sub(r"\s+\d+$", "", linha)
    return linha.strip()


def detectar_nivel(texto: str) -> int:
    if texto.lower().startswith(("refer", "apêndic", "anex")):
        return 1
    match = re.match(r"^(\d+(?:\.\d+)*)\s+", texto)
    if match:
        return len(match.group(1).split("."))
    return 1


def formatar_run(run, nivel: int, texto: str) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = PRETO
    if nivel <= 3:
        run.bold = True
    if nivel == 1 and not texto.lower().startswith("refer"):
        run.text = run.text.upper()


def parse_prefix(text: str):
    if text.strip().upper().startswith("REFER"):
        return "REF"
    if text.strip().upper().startswith("APÊNDIC"):
        return "APE"
    if text.strip().upper().startswith("ANEX"):
        return "ANE"
    match = re.match(r"^(\d+(?:\.\d+)*)\s+", text.strip())
    if match:
        return match.group(1)
    return None


def processar_substituicao(doc_path: Path, txt_path: Path, out_path: Path) -> None:
    print(f"Lendo documento: {doc_path}")
    doc = Document(str(doc_path))

    placeholders = []
    known_ph_texts = [
        "1 INTRODUÇÃO", "2 DESENVOLVIMENTO",
        "2.1 Tópico secundário", "2.1.1 Tópico terciário",
        "2.1.1.1 Tópico quaternário", "2.1.1.1.1 Tópico quinário",
        "3 TÓPICO PRIMÁRIO DO DESENVOLVIMENTO", "4 RESULTADOS E DISCUSSÕES",
        "5 CONSIDERAÇÕES FINAIS", "REFERÊNCIAS", "APÊNDICES", "ANEXOS",
    ]
    known_upper = {k.upper() for k in known_ph_texts}

    for p in doc.paragraphs:
        ptxt = p.text.strip().upper()
        pfx = parse_prefix(p.text)
        if ptxt in known_upper and pfx:
            placeholders.append({"prefix": pfx, "paragraph": p, "used": False})

    with open(str(txt_path), "r", encoding="utf-8-sig") as f:
        linhas = f.readlines()

    topicos_novos = []
    for linha in linhas:
        texto = limpar_texto_topico(linha)
        if texto:
            topicos_novos.append({"prefix": parse_prefix(texto), "texto": texto, "nivel": detectar_nivel(texto)})

    def get_ph(pfx):
        for ph in placeholders:
            if ph["prefix"] == pfx:
                return ph
        return None

    def get_next_ph_paragraph(current_new_idx):
        for i in range(current_new_idx + 1, len(topicos_novos)):
            ph = get_ph(topicos_novos[i]["prefix"])
            if ph:
                return ph["paragraph"]
        ref_ph = get_ph("REF")
        if ref_ph:
            return ref_ph["paragraph"]
        return None

    for i, topico in enumerate(topicos_novos):
        pfx = topico["prefix"]
        texto = topico["texto"]
        nivel = topico["nivel"]

        ph = get_ph(pfx)
        if ph:
            p = ph["paragraph"]
            p.clear()
            run = p.add_run(texto)
            formatar_run(run, nivel, texto)
            ph["used"] = True
        else:
            target_p = get_next_ph_paragraph(i)
            if target_p:
                p = target_p.insert_paragraph_before()
            else:
                p = doc.add_paragraph()

            style_name = f"Heading {nivel}" if nivel <= 9 else "Heading 1"
            try:
                p.style = style_name
            except Exception:
                pass
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(texto)
            formatar_run(run, nivel, texto)

    for ph in placeholders:
        if not ph["used"] and ph["prefix"] not in ("REF", "APE", "ANE"):
            p = ph["paragraph"]
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Substituição concluída! Arquivo salvo em: {out_path}")


def main() -> None:
    if len(sys.argv) < 4:
        print("Uso: python -m scripts.gerenciador_topicos <arquivo.docx> <topicos.txt> <saida.docx>")
        sys.exit(1)
    processar_substituicao(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))


if __name__ == "__main__":
    main()
